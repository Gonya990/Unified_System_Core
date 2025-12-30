from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from docx.oxml import OxmlElement

def set_rtl(paragraph):
    """Sets the paragraph direction to RTL."""
    try:
        paragraph.paragraph_format.bidi = True
    except AttributeError:
        # Fallback if property doesn't exist (unlikely in 1.2.0)
        pPr = paragraph._p.get_or_add_pPr()
        from docx.oxml import OxmlElement
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

def create_resume():
    document = Document()
    
    # Define Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # ---------------- Header ----------------
    # Name
    name_paragraph = document.add_paragraph()
    set_rtl(name_paragraph)
    name_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = name_paragraph.add_run("איגור גונצ׳רנקו")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(44, 62, 80) # Dark Blue
    
    # Subtitle
    title_paragraph = document.add_paragraph()
    set_rtl(title_paragraph)
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = title_paragraph.add_run("מנהל תפעול שטח | אינטגרטור טכני")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(85, 85, 85)
    run.bold = True

    # Contact Info
    contact_paragraph = document.add_paragraph()
    set_rtl(contact_paragraph)
    contact_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    contact_paragraph.add_run("[Email] | [Phone] | מרטין בובר 7, קרית ביאליק")
    
    document.add_heading('תמצית מקצועית', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    summary_p = document.add_paragraph()
    set_rtl(summary_p)
    summary_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    summary_p.add_run("מנהל טכני ומנהל עבודה בעל ניסיון עשיר בהקמה ותחזוקת תשתיות מורכבות בתחומי המכאניקה, החשמל והתקשורת. בעל יכולת מוכחת בהובלת צוותים טכניים וביצוע פרויקטים מורכבים בסביבות עבודה תובעניות, תוך הקפדה יתרה על נהלי בטיחות ואיכות. משלב רקע טכני \"Hands-on\" עם יכולות למידה עצמית גבוהות ויוזמה, הבאות לידי ביטוי בשליטה בטכנולוגיות מתקדמות, אוטומציה ופתרונות אינטגרציה.")

    # ---------------- Experience ----------------
    document.add_heading('ניסיון תעסוקתי', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Job 1
    job1_header = document.add_paragraph()
    set_rtl(job1_header)
    job1_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job1_header.add_run("Schindler Group | סופרוויזור (BEYOND Tower)")
    run.bold = True
    run.font.size = Pt(12)
    
    job1_meta = document.add_paragraph()
    set_rtl(job1_meta)
    job1_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job1_meta.add_run("אוגוסט 2023 – הווה | ישראל")
    run.italic = True
    run.font.color.rgb = RGBColor(127, 140, 141)

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("פיקוח טכני: ").bold = True
    p.add_run("ניהול והובלת התקנת מערכות מעליות מהירות בפרויקטי דגל רבי-קומות. אחריות כוללת על פתרון בעיות מכאניות ולוגיסטיות בשטח בזמן אמת.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("בטיחות ותקינה: ").bold = True
    p.add_run("יישום והטמעה של נהלי בטיחות מחמירים בעבודה בגובה ועמידה בסטנדרטים בינלאומיים של איכות.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("ניהול ודיווח: ").bold = True
    p.add_run("מעקב שוטף אחר התקדמות הפרויקט, ניהול יומני עבודה ודיווח להנהלת הפרויקט באמצעות מערכות דיגיטליות.")

    # Job 2
    document.add_paragraph() # Spacer
    job2_header = document.add_paragraph()
    set_rtl(job2_header)
    job2_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job2_header.add_run("URICOMS (Partner Communications) | מומחה סיבים אופטיים / ראש צוות")
    run.bold = True
    run.font.size = Pt(12)
    
    job2_meta = document.add_paragraph()
    set_rtl(job2_meta)
    job2_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job2_meta.add_run("2019 – 2023 | ישראל")
    run.italic = True
    run.font.color.rgb = RGBColor(127, 140, 141)

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("תשתיות רשת: ").bold = True
    p.add_run("התמחות בפריסת תשתיות נתונים רחבות פס (FTTx), ביצוע ריתוכי סיבים מדויקים ובדיקות תשתית (OTDR). ניהול תהליך ההתקנה מקצה לקצה.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("הובלה טכנית: ").bold = True
    p.add_run("הדרכה וחניכה של צוותי שטח, הטמעת שיטות עבודה לניהול כבילה יעיל ופתרון תקלות רשת מורכבות בבית הלקוח ובתשתיות הבניין.")

    # Job 3
    document.add_paragraph() # Spacer
    job3_header = document.add_paragraph()
    set_rtl(job3_header)
    job3_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job3_header.add_run("וירה והרכבות בע״מ | טכנאי מעליות")
    run.bold = True
    run.font.size = Pt(12)
    
    job3_meta = document.add_paragraph()
    set_rtl(job3_meta)
    job3_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = job3_meta.add_run("2020 – 2022 | ישראל")
    run.italic = True
    run.font.color.rgb = RGBColor(127, 140, 141)

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("התקנה והרכבה: ").bold = True
    p.add_run("ביצוע עבודות מכאניות מורכבות הכוללות הרכבת מסילות, מנועים ורכיבי תא.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("קריאת שרטוטים: ").bold = True
    p.add_run("עבודה צמודה עם שרטוטים טכניים והוראות יצרן לביצוע התקנות מדויקות.")


    # ---------------- Skills ----------------
    document.add_heading('כישורים טכניים ופרויקטים', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Skills Sub-header
    sk_h = document.add_paragraph()
    set_rtl(sk_h)
    sk_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sk_h.add_run("מיומנויות מקצועיות").bold = True

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("עבודת שטח: ").bold = True
    p.add_run("התקנות מכאניות, חיווט מתח נמוך, סיבים אופטיים, ניהול לוגיסטי בשטח.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("כלים ותוכנות: ").bold = True
    p.add_run("שליטה מלאה ביישומי Office, מערכות לניהול פרויקטים, ויכולת קריאה והבנה של שרטוטים טכניים.")
    
    document.add_paragraph() # Spacer

    # Project Sub-header
    pr_h = document.add_paragraph()
    set_rtl(pr_h)
    pr_h.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = pr_h.add_run('פרויקט אישי: "תשתיות חכמות ואוטומציה"')
    run.bold = True
    
    pr_desc = document.add_paragraph()
    set_rtl(pr_desc)
    pr_desc.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pr_desc.add_run("תכנון, הקמה ותחזוקה של מעבדה ביתית מתקדמת, המדגימה יכולות אינטגרציה ולמידה טכנולוגית:")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("בית חכם ו-IoT: ").bold = True
    p.add_run("אינטגרציה של פרוטוקולים שונים מול שרת Home Assistant מרכזי לאוטומציה מלאה.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("מערכות הפעלה ו-AI: ").bold = True
    p.add_run("ניהול שרתי Linux וסביבות וירטואליות (Proxmox/Docker) להרצת מודלי AI מקומיים.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("פתרון בעיות: ").bold = True
    p.add_run("שימוש מתקדם בכלי AI לכתיבת סקריפטים ופתרונות ניטור מותאמים אישית.")

    # ---------------- Education ----------------
    document.add_heading('השכלה ורישיונות', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("הכשרה במכונאות ").bold = True
    p.add_run("(לימודים בהיקף 3 שנים)")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("רישיונות והסמכות: ").bold = True
    p.add_run("מלגזה (20 טון), רישיון נהיגה ב', אישור לעבודה בגובה.")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("שירות צבאי: ").bold = True
    p.add_run("חיל האוויר, תפקיד נהג, דרגת סמ״ר (2013-2015).")

    # ---------------- Languages ----------------
    document.add_heading('שפות', level=1).alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("עברית: ").bold = True
    p.add_run("שליטה טובה מאוד")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("רוסית: ").bold = True
    p.add_run("שפת אם")

    p = document.add_paragraph(style='List Bullet')
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("אנגלית: ").bold = True
    p.add_run("טכנית בסיסית")

    # Save
    document.save('Resume_Igor_Refined_✅.docx')

if __name__ == "__main__":
    create_resume()
