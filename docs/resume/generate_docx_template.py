from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# --- Helpers ---

def set_rtl(paragraph):
    """Sets the paragraph direction to RTL."""
    try:
        paragraph.paragraph_format.bidi = True
    except AttributeError:
        pass
    # OXML fallback for robustness
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

def set_cell_background(cell, color_hex):
    """Set background shading for a table cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Check if shd already exists to avoid duplication errors (simplified)
    # Ideally remove existing, but for new doc it's fine.
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_header_bar(paragraph, text, bg_color="1B4D3E"): # Dark Green
    """Creates a text block with a background color (shading) for headers."""
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.size = Pt(14)
    # Apply shading to the run or paragraph? Paragraph shading is better for full width bars
    # But python-docx paragraph shading limits. We will simulate with paragraph border or shading.
    # Actually, let's use the paragraph shading OXML.
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), bg_color)
    pPr.append(shd)

    # Add some spacing
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)

def create_resume_template():
    document = Document()

    # 1. Set Margin to narrow to maximize space
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # 2. Styles
    style = document.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    # 3. Create Main Layout Table (1 Row, 2 Cols)
    table = document.add_table(rows=1, cols=2)
    table.autofit = False

    # Set Table Direction to RTL
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    bidiViz = OxmlElement('w:bidiVisual')
    tblPr.append(bidiViz)

    # Column Widths
    # Total width approx 18cm.
    # Sidebar (Right in RTL) -> ~30% -> 6cm
    # Main (Left in RTL) -> ~70% -> 12cm

    # In RTL table: Cell 0 is Right, Cell 1 is Left.
    sidebar_cell = table.rows[0].cells[0]
    main_cell = table.rows[0].cells[1]

    sidebar_cell.width = Cm(6.5)
    main_cell.width = Cm(12.5)

    # ---------------- SIDEBAR (Right) ----------------
    set_cell_background(sidebar_cell, "89CFF0") # Baby Blue / Light Blue

    # Placeholder for Photo (Circle text)
    p_photo = sidebar_cell.paragraphs[0]
    set_rtl(p_photo)
    p_photo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_photo.paragraph_format.space_after = Pt(20)
    run_photo = p_photo.add_run("[תמונה]")
    run_photo.font.size = Pt(10)
    run_photo.font.color.rgb = RGBColor(100, 100, 100)

    # Name
    p_name = sidebar_cell.add_paragraph()
    set_rtl(p_name)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("איגור\nגונצ׳רנקו")
    run_name.bold = True
    run_name.font.size = Pt(28)
    run_name.font.color.rgb = RGBColor(40, 40, 40)

    # Title
    p_title = sidebar_cell.add_paragraph()
    set_rtl(p_title)
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("מנהל תפעול שטח\nאינטגרטור טכני")
    run_title.font.size = Pt(14)
    run_title.font.color.rgb = RGBColor(255, 255, 255) # White text for contrast often looks good on blue, or dark gray? Template has dark text.
    run_title.font.color.rgb = RGBColor(60, 60, 60)
    p_title.paragraph_format.space_after = Pt(30)

    # Separator
    # sidebar_cell.add_paragraph("--------------------").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact Info
    p_contact_head = sidebar_cell.add_paragraph()
    set_rtl(p_contact_head)
    p_contact_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_ch = p_contact_head.add_run("פרטי קשר")
    run_ch.bold = True
    run_ch.font.size = Pt(14)
    run_ch.font.color.rgb = RGBColor(255, 255, 255) # White headers in sidebar often pop

    p_contact = sidebar_cell.add_paragraph()
    set_rtl(p_contact)
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = "מרטין בובר 7, קרית ביאליק\n[Phone]\n[Email]"
    run_c = p_contact.add_run(contact_text)
    run_c.font.size = Pt(10)
    run_c.font.color.rgb = RGBColor(50, 50, 50)
    p_contact.paragraph_format.space_after = Pt(20)

    # Languages
    p_lang_head = sidebar_cell.add_paragraph()
    set_rtl(p_lang_head)
    p_lang_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_lh = p_lang_head.add_run("שפות")
    run_lh.bold = True
    run_lh.font.size = Pt(14)
    run_lh.font.color.rgb = RGBColor(255, 255, 255)

    p_lang = sidebar_cell.add_paragraph()
    set_rtl(p_lang)
    p_lang.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lang_text = "עברית\n(שליטה טובה מאוד)\n\nרוסית\n(שפת אם)\n\nאנגלית\n(טכנית בסיסית)"
    run_l = p_lang.add_run(lang_text)
    run_l.font.color.rgb = RGBColor(50, 50, 50)
    p_lang.paragraph_format.space_after = Pt(20)

    # Skills Summary for Sidebar
    p_skill_head = sidebar_cell.add_paragraph()
    set_rtl(p_skill_head)
    p_skill_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sh = p_skill_head.add_run("כישורים")
    run_sh.bold = True
    run_sh.font.size = Pt(14)
    run_sh.font.color.rgb = RGBColor(255, 255, 255)

    p_skills = sidebar_cell.add_paragraph()
    set_rtl(p_skills)
    p_skills.alignment = WD_ALIGN_PARAGRAPH.CENTER
    skills_text = "ניהול צוותים\nפיקוח טכני\nסיבים אופטיים\nמערכות מעליות\nבטיחות וגה\"צ\nOffice & ERP"
    run_s = p_skills.add_run(skills_text)
    run_s.font.color.rgb = RGBColor(50, 50, 50)


    # ---------------- MAIN CONTENT (Left) ----------------
    set_cell_background(main_cell, "FAF9F6") # Off White / Ivory

    # Summary Section
    p_sum_head = main_cell.paragraphs[0]
    set_rtl(p_sum_head)
    create_header_bar(p_sum_head, "   תמצית מקצועית   ")

    p_summary = main_cell.add_paragraph()
    set_rtl(p_summary)
    p_summary.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_summary.add_run("מנהל טכני ומנהל עבודה בעל ניסיון עשיר בהקמה ותחזוקת תשתיות מורכבות בתחומי המכאניקה, החשמל והתקשורת. בעל יכולת מוכחת בהובלת צוותים טכניים וביצוע פרויקטים מורכבים בסביבות עבודה תובעניות, תוך הקפדה יתרה על נהלי בטיחות ואיכות. משלב רקע טכני \"Hands-on\" עם יכולות למידה עצמית גבוהות ויוזמה.")
    p_summary.paragraph_format.space_after = Pt(15)

    # Experience Section
    p_exp_head = main_cell.add_paragraph()
    set_rtl(p_exp_head)
    create_header_bar(p_exp_head, "   ניסיון תעסוקתי   ")

    # Job 1
    p_job1 = main_cell.add_paragraph()
    set_rtl(p_job1)
    p_job1.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    r_date1 = p_job1.add_run("2023 – הווה")
    r_date1.bold = True
    r_date1.font.size = Pt(11)
    p_job1.add_run("\n")
    r_role1 = p_job1.add_run("סופרוויזור (BEYOND Tower)")
    r_role1.bold = True
    r_role1.font.color.rgb = RGBColor(0, 100, 200) # Blueish title like in template
    r_role1.font.size = Pt(12)
    p_job1.add_run(" | ")
    r_comp1 = p_job1.add_run("Schindler Group")
    r_comp1.bold = False
    r_comp1.font.color.rgb = RGBColor(100, 100, 100)

    p_desc1 = main_cell.add_paragraph()
    set_rtl(p_desc1)
    p_desc1.add_run("• ניהול והובלת התקנת מערכות מעליות מהירות בפרויקטי דגל רבי-קומות.\n• אחריות כוללת על פתרון בעיות מכאניות ולוגיסטיות בשטח בזמן אמת.\n• יישום נהלי בטיחות ומעקב אחר התקדמות הפרויקט.")
    p_desc1.paragraph_format.space_after = Pt(12)

    # Job 2
    p_job2 = main_cell.add_paragraph()
    set_rtl(p_job2)

    r_date2 = p_job2.add_run("2019 – 2023")
    r_date2.bold = True
    p_job2.add_run("\n")
    r_role2 = p_job2.add_run("מומחה סיבים אופטיים / ראש צוות")
    r_role2.bold = True
    r_role2.font.color.rgb = RGBColor(0, 100, 200)
    r_role2.font.size = Pt(12)
    p_job2.add_run(" | ")
    r_comp2 = p_job2.add_run("URICOMS (Partner)")
    r_comp2.font.color.rgb = RGBColor(100, 100, 100)

    p_desc2 = main_cell.add_paragraph()
    set_rtl(p_desc2)
    p_desc2.add_run("• התמחות בפריסת תשתיות FTTx וביצוע ריתוכי סיבים.\n• ניהול תהליך ההתקנה מקצה לקצה והדרכת צוותי שטח.")
    p_desc2.paragraph_format.space_after = Pt(12)

    # Job 3
    p_job3 = main_cell.add_paragraph()
    set_rtl(p_job3)

    r_date3 = p_job3.add_run("2020 – 2022")
    r_date3.bold = True
    p_job3.add_run("\n")
    r_role3 = p_job3.add_run("טכנאי מעליות")
    r_role3.bold = True
    r_role3.font.color.rgb = RGBColor(0, 100, 200)
    r_role3.font.size = Pt(12)
    p_job3.add_run(" | ")
    r_comp3 = p_job3.add_run("וירה והרכבות בע״מ")
    r_comp3.font.color.rgb = RGBColor(100, 100, 100)

    p_desc3 = main_cell.add_paragraph()
    set_rtl(p_desc3)
    p_desc3.add_run("• הרכבת מסילות, מנועים ורכיבי תא.\n• עבודה עם שרטוטים טכניים.")
    p_desc3.paragraph_format.space_after = Pt(15)

    # Education Section
    p_edu_head = main_cell.add_paragraph()
    set_rtl(p_edu_head)
    create_header_bar(p_edu_head, "   השכלה ורישיונות   ")

    p_edu = main_cell.add_paragraph()
    set_rtl(p_edu)
    p_edu.add_run("הכשרה במכונאות").bold = True
    p_edu.add_run("\nלימודים בהיקף 3 שנים")
    p_edu.add_run("\n\n")
    p_edu.add_run("רישיונות:").bold = True
    p_edu.add_run(" מלגזה (20 טון), נהיגה ב', עבודה בגובה.")

    # Military Section
    p_mil_head = main_cell.add_paragraph()
    set_rtl(p_mil_head)
    create_header_bar(p_mil_head, "   שירות צבאי   ")

    p_mil = main_cell.add_paragraph()
    set_rtl(p_mil)
    p_mil.add_run("2013 – 2015").bold = True
    p_mil.add_run("\nחיל האוויר, נהג, דרגת סמ״ר.")

    # Project (Optional, maybe fit in if space allows, or leave out to keep it clean like template)
    p_proj_head = main_cell.add_paragraph()
    set_rtl(p_proj_head)
    create_header_bar(p_proj_head, "   פרויקטים אישיים   ")

    p_proj = main_cell.add_paragraph()
    set_rtl(p_proj)
    p_proj.add_run("תשתיות חכמות ואוטומציה: ").bold = True
    p_proj.add_run("הקמת מעבדה ביתית (Home Assistant, Linux, AI) - מדגים יכולת למידה עצמית ואינטגרציה.")


    # Save
    document.save('Resume_Igor_Refined_Template_✅.docx')

if __name__ == "__main__":
    create_resume_template()
