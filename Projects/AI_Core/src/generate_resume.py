from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_resume():
    document = Document()

    # --- HEADER ---
    name = document.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name.add_run("איגור גונצ׳רנקו")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Arial"

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("מנהל תפעול | מנהל פרויקטים טכני | אינטגרטור")
    run.font.size = Pt(14)
    run.font.name = "Arial"

    # Contacts
    contacts = document.add_paragraph()
    contacts.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contacts.add_run("Email: gonya90.gg@gmail.com | Mobile: 053-XXX-XXXX | קרית ביאליק\n")
    contacts.add_run("LinkedIn: linkedin.com/in/igor-goncharenko | Telegram: @igoreha9")

    document.add_paragraph()  # Spacer

    # --- Helper function for Sections ---
    def add_section_header(text):
        p = document.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.underline = True
        p.space_after = Pt(6)

    # --- PROFILE ---
    add_section_header("תמצית מנהלים (PROFILE)")
    p = document.add_paragraph(
        "מנהל תפעול ופרויקטים מנוסה, המתמחה בהובלת צוותים טכניים, ניהול לוגיסטיקה מורכבת ופיקוח על פרויקטי תשתיות (בנייה לגובה, מערכות אלקטרו-מכאניות, תקשורת).\n"
        "בעל יכולת מוכחת בפתרון בעיות בשטח בזמן אמת (Troubleshooting), ניהול משאבים ואינטגרציה בין מערכות טכנולוגיות.\n"
        "משלב רקע טכני חזק (Hands-on) עם ראייה ניהולית רחבה, תודעת שירות גבוהה ויכולות למידה עצמית (אוטודידקט) בתחומי האוטומציה וה-AI.\n"
        "מחפש תפקיד מאתגר בניהול תפעול / לוגיסטיקה / פרויקטים / פיקוח טכני."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- EXPERIENCE ---
    add_section_header("ניסיון תעסוקתי (PROFESSIONAL EXPERIENCE)")

    # Schindler
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("2023 – היום | Schindler Group | מפקח טכני / מנהל עבודה (Supervisor - Major Projects)")
    r.bold = True

    items = [
        "ניהול ופיקוח כולל על פרויקטי הקמת מעליות במגדלי יוקרה (High-rise), משלב התכנון ועד המסירה.",
        "ניהול אופרטיבי של אתר העבודה: תיאום מול קבלנים, הנדסה, בטיחות ולוחות זמנים (Gantt).",
        "אחריות לוגיסטית: ניהול מלאי אתר, הזמנת ציוד, קבלת סחורה וניהול שרשרת אספקה.",
        "אינטגרציה: חיבור מערכות פיקוד ובקרה, פתרון תקלות טכניות מורכבות בשטח.",
        "בטיחות ואיכות: הטמעת נהלי בטיחות גלובליים, בקרת איכות (QA/QC) והדרכת עובדים.",
    ]
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Partner
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("2019 – 2023 | URICOMS (עבור Partner) | ראש צוות תשתיות וסיבים (Team Leader)")
    r.bold = True

    items = [
        "ניהול והובלת צוותי טכנאים בפריסת תשתיות סיבים אופטיים (FTTH) ותקשורת.",
        "אחריות על עמידה ביעדי ביצוע (KPIs / SLA), איכות התקנה ושביעות רצון לקוחות.",
        "הדרכה וחניכה מקצועית של עובדים, הטמעת נהלי עבודה ושיפור תהליכים.",
        "עבודה בסביבה דינמית תחת לחץ, מתן פתרונות טכניים לבעיות תשתית בבניינים ועסקים.",
    ]
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Vira
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run("2020 – 2022 | וירה והרכבות בע״מ | טכנאי מעליות ודרגנועים")
    r.bold = True

    items = [
        "ביצוע התקנות מכאניות וחשמליות (Heavy Duty), קריאת שרטוטים הנדסיים.",
        "עבודה בצוותים, הקפדה על דיוק ואיכות ביצוע ברמה גבוהה.",
    ]
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- EDUCATION ---
    add_section_header("השכלה והסמכות (EDUCATION & SKILLS)")
    items = [
        "הנדסאי/טכנאי מכונות (לימודי תעודה, 3 שנים).",
        "רישיונות: רישיון נהיגה ב', רישיון מלגזה (20 טון), הסמכה לעבודה בגובה.",
        "שירות צבאי: חיל האוויר, נהג מבצעי, דרגת סמ״ר.",
    ]
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # --- LANG & TECH ---
    add_section_header("שפות וטכנולוגיה (LANGUAGES & TECH)")
    p = document.add_paragraph()
    p.add_run("• עברית: רמה גבוהה מאוד / שפת עבודה.\n")
    p.add_run("• רוסית: שפת אם.\n")
    p.add_run("• אנגלית: טכנית טובה.\n")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    add_section_header("מיומנויות נוספות (ADDITIONAL SKILLS)")
    items = [
        "ניהול מחסן ולוגיסטיקה ממוחשב, ERP, Office (Excel/Outlook).",
        "טכנולוגיה מתקדמת: הקמת שרתי בית חכם (Home Assistant, MQTT), לינוקס (Linux), וירטואליזציה (Proxmox), שימוש בכלי AI לשיפור פרודוקטיביות.",
        "תכונות: אמינות, דייקנות, יחסי אנוש מעולים, יכולת הנעת עובדים.",
    ]
    for item in items:
        p = document.add_paragraph(item, style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    output_path = "/Users/macbook/Desktop/Resume_Igor_Perfect_✅.docx"
    document.save(output_path)
    print(f"Resume saved to {output_path}")


if __name__ == "__main__":
    create_resume()
